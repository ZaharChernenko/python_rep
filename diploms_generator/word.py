import json
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


def fill_template(json_file, template_path, output_dir):

    # Загрузка данных из JSON
    with open(json_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document(template_path)

    # Определяем слово на основе пола
    if data["Пол"] == "м":
        слово_выступивший = "выступившему"
    elif data["Пол"] == "ж":
        слово_выступивший = "выступившей"

    # Словарь замен
    replacements = {
        "{ФИО}": (data["ФИО в Дательном падеже"], Pt(55), RGBColor(112, 48, 160), True, True),
        "{Курс}": (str(data["Курс"]), Pt(14), None, False, False),
        "{Группа}": (data["Группа"], Pt(14), None, False, False),
        "{Номинация}": (data["Номинация"], Pt(14), None, True, False),
        "{ТемаДоклада}": (data["Тема доклада"], Pt(16), None, False, True),
        "{НаучныйРуководитель}": (data["Научный руководитель"], None, None, False, False),
        "{Дата}": (data["Дата"], Pt(16), None, False, True),
        "{НомерТура}": (data["Номер тура в Родительном падеже"], Pt(16), None, False, False),
        "{Выступивший}": (слово_выступивший, Pt(14), None, False, False),
    }

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)

    # Замена текста с форматированием
    for paragraph in doc.paragraphs:
        for placeholder, (replacement, font_size, font_color, bold, italic) in replacements.items():
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
                    if font_size:
                        run.font.size = font_size
                    if font_color:
                        run.font.color.rgb = font_color
                    run.bold = bold
                    run.italic = italic

    # Сохранение документа
    word_output_path = os.path.join(output_dir, f'{data["ФИО в Дательном падеже"]}.docx')
    doc.save(word_output_path)

    print(f"Сертификат для {data['ФИО в Дательном падеже']} сохранён в {word_output_path}")


# Пути
json_path = "./data.json"
template_path = "./Сертификат.docx"
output_directory = "./output"

os.makedirs(output_directory, exist_ok=True)
fill_template(json_path, template_path, output_directory)
