from typing import NamedTuple

from docx import Document
from docx.enum.text import WD_LINE_SPACING

from .pattern_builder import TextData


class WordFillerOutput(NamedTuple):
    document: Document
    file_name: str


class WordFiller:
    def __init__(self, template_file_path: str):
        self.template_file_path: str = template_file_path

    def __call__(self, pattern_obj: dict[str, TextData]) -> WordFillerOutput:
        doc = Document(self.template_file_path)

        for paragraph in doc.paragraphs:
            full_text = "".join(run.text for run in paragraph.runs)
            modified_text = full_text

            for placeholder, text_data in pattern_obj.items():
                if placeholder in full_text:
                    modified_text = modified_text.replace(placeholder, text_data.text)

            if full_text != modified_text:
                for run in paragraph.runs:
                    run.text = ""

                new_run = paragraph.add_run(modified_text)

                for placeholder, text_data in pattern_obj.items():
                    if text_data.text in new_run.text:
                        new_run.font.size = text_data.font_size
                        new_run.font.color.rgb = text_data.font_color
                        new_run.bold = text_data.is_bold
                        new_run.italic = text_data.is_italic

                        if text_data.line_spacing is not None:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.line_spacing = text_data.line_spacing
                            paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        return WordFillerOutput(
            document=doc, file_name=f"{pattern_obj['{Фамилия}'].text}{pattern_obj['{Имя}'].text}.docx"
        )
